#!/usr/bin/env python3
"""
Generate DD-Grade Audit Executive Summary for TrendStudio Investor Model v1.0.2 Public.
Output: audit_public_v1_executive_summary.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

H1_COLOR = RGBColor(0x00, 0x70, 0xC0)
H2_COLOR = RGBColor(0x00, 0x70, 0xC0)
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
H1_SIZE = Pt(22)
H2_SIZE = Pt(16)

RED = RGBColor(0xCC, 0x00, 0x00)
AMBER = RGBColor(0xFF, 0x8C, 0x00)
GREEN = RGBColor(0x22, 0x8B, 0x22)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run.font.size = H1_SIZE
        run.font.color.rgb = H1_COLOR
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run.font.size = H2_SIZE
        run.font.color.rgb = H2_COLOR
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = size or BODY_SIZE
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    return p


def add_mixed_para(doc, fragments):
    """Add a paragraph with mixed formatting.
    fragments: list of (text, bold, italic, color, size) tuples.
    """
    p = doc.add_paragraph()
    for text, bold, italic, color, size in fragments:
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run.font.size = size or BODY_SIZE
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix="", color=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = BODY_FONT
        run_b._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_b.font.size = BODY_SIZE
        if color:
            run_b.font.color.rgb = color
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = BODY_SIZE
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "0070C0")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = BODY_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            run.font.size = Pt(11)
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8F0FE")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ---------------------------------------------------------------------------
# Main document generation
# ---------------------------------------------------------------------------

def generate():
    doc = Document()

    # -- Page setup: A4, margins 3/1.5/2/2 cm (top/bottom/left/right) --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # -- Default style --
    style = doc.styles["Normal"]
    font = style.font
    font.name = BODY_FONT
    font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    # ======================================================================
    # TITLE BLOCK
    # ======================================================================
    add_heading1(doc, "Audit Executive Summary")
    add_heading2(doc, "Investor Model v1.0.2 Public \u2014 DD-Grade Audit")

    add_mixed_para(doc, [
        ("Date: ", True, False, None, BODY_SIZE),
        ("April 13, 2026", False, False, None, BODY_SIZE),
        ("  |  ", False, False, None, BODY_SIZE),
        ("Auditor: ", True, False, None, BODY_SIZE),
        ("Independent (Claude Code)", False, False, None, BODY_SIZE),
        ("  |  ", False, False, None, BODY_SIZE),
        ("Scope: ", True, False, None, BODY_SIZE),
        ("7 blocks (A\u2013G)", False, False, None, BODY_SIZE),
    ])

    # Thin horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(2)
    p_hr.paragraph_format.space_after = Pt(8)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="0070C0"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ======================================================================
    # 1. OVERALL VERDICT
    # ======================================================================
    add_heading2(doc, "1. Overall Verdict")

    add_mixed_para(doc, [
        ("CONDITIONAL FAIL", True, False, RED, Pt(16)),
        (" \u2014 Model requires remediation before investor distribution.", False, False, None, BODY_SIZE),
    ])

    add_para(doc, "Traffic Light Summary:", bold=True)

    traffic_light = [
        ("Block A (xlsx Technical): ", "AMBER", AMBER,
         " \u2014 structurally sound OOXML, but 97.5% static cells, metadata leaks"),
        ("Block B (Financial Math): ", "AMBER", AMBER,
         " \u2014 anchors verified (Revenue 4,545, NDP 3,000, IRR 20.09%), but back-solved architecture, "
         "3 inconsistent IRR methods, MC mean IRR 11.44% (8.65pp below deterministic)"),
        ("Block C (Public\u2194Internal): ", "RED", RED,
         " \u2014 Critical leakage of Internal waterfall W5 V\u2011D comparison, "
         "8 \u201cInternal\u201d keyword occurrences, file path exposure"),
        ("Block D (Investor Readiness): ", "AMBER", AMBER,
         " \u2014 Stress tests show resilience (cash positive through 2032, BS balanced), "
         "but 6\u00d7 valuation spread, P(IRR>hurdle) = 13.6%, fabricated comps"),
        ("Block E (Going Concern): ", "GREEN", GREEN,
         " \u2014 Cash end 2032 = 891.75M, BS balanced all periods, no covenant breach in base case"),
        ("Block F (Cover Docs): ", "AMBER", AMBER,
         " \u2014 Comprehensive manifest and honest self-audit (34 known issues documented), "
         "but template placeholders remain, cover letter incomplete"),
        ("Block G (Pipeline): ", "GREEN", GREEN,
         " \u2014 Impressive 16,621-line pipeline with 287 tests, Pydantic validation, "
         "seed-based determinism, 8 ADRs. Minor: unpinned requirements, 3 hardcoded paths."),
    ]

    for prefix, status, status_color, detail in traffic_light:
        p = doc.add_paragraph(style="List Bullet")
        run_prefix = p.add_run(prefix)
        run_prefix.bold = True
        run_prefix.font.name = BODY_FONT
        run_prefix._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_prefix.font.size = BODY_SIZE

        run_status = p.add_run(status)
        run_status.bold = True
        run_status.font.name = BODY_FONT
        run_status._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_status.font.size = BODY_SIZE
        run_status.font.color.rgb = status_color

        run_detail = p.add_run(detail)
        run_detail.font.name = BODY_FONT
        run_detail._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_detail.font.size = BODY_SIZE
        p.paragraph_format.space_after = Pt(2)

    # ======================================================================
    # 2. TOP 5 CRITICAL FINDINGS
    # ======================================================================
    add_heading2(doc, "2. Top 5 Critical Findings")

    findings = [
        ("1. INTERNAL DATA LEAKAGE (P0): ",
         "24_Investor_Returns!B49 explicitly exposes W5 V\u2011D waterfall comparison text. "
         "8 \u201cInternal\u201d references in sharedStrings. Developer path in workbook.xml."),
        ("2. STATIC MODEL (P1): ",
         "226 formulas / 9,000+ values = 2.5%. Not a live financial model. Will fail DD instantly."),
        ("3. VALUATION SPREAD 6\u00d7 (P1): ",
         "DCF 1.8B vs Comps 7.5B vs MC 11.2B. Unexplained gap kills credibility."),
        ("4. MC vs DETERMINISTIC GAP (P1): ",
         "Base IRR 20.09% vs MC Mean 11.44%. P(IRR>18% hurdle) = 13.6%."),
        ("5. THREE IRR METHODS (P1): ",
         "Newton\u2019s, MOIC approximation, numpy \u2014 produce different results for same inputs."),
    ]

    for prefix, detail in findings:
        p = doc.add_paragraph()
        run_b = p.add_run(prefix)
        run_b.bold = True
        run_b.font.name = BODY_FONT
        run_b._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_b.font.size = BODY_SIZE
        run_b.font.color.rgb = RED

        run_d = p.add_run(detail)
        run_d.font.name = BODY_FONT
        run_d._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_d.font.size = BODY_SIZE
        p.paragraph_format.space_after = Pt(4)

    # ======================================================================
    # 3. KEY FINANCIAL VERIFICATION
    # ======================================================================
    add_heading2(doc, "3. Key Financial Verification")

    add_para(doc, "Anchor values verified from xlsx:", italic=True)

    fin_headers = ["Metric", "Model Value", "Expected", "Status"]
    fin_rows = [
        ["Revenue 3Y", "4,545", "4,545", "PASS"],
        ["EBITDA GAAP 3Y", "2,167.4", "2,076.1 (prompt)", "MISMATCH \u2014 91.3 delta"],
        ["NDP", "3,000", "3,000", "PASS"],
        ["IRR W3 Base", "20.09%", "20.09%", "PASS"],
        ["MC Mean IRR", "11.44%", "11.44%", "PASS"],
        ["BS Balance", "0.00 all periods", "0", "PASS"],
        ["Cash End 2032", "891.75", ">0", "PASS"],
    ]
    make_table(doc, fin_headers, fin_rows, col_widths=[4.5, 3.5, 4.0, 5.0])

    doc.add_paragraph()  # spacer

    add_mixed_para(doc, [
        ("Note: ", True, False, AMBER, BODY_SIZE),
        ("EBITDA discrepancy \u2014 prompt states 2,076.1 but actual model shows 2,167.4 "
         "(post-FOT A2 cascade v1.0.2).", False, True, None, BODY_SIZE),
    ])

    # ======================================================================
    # 4. GOING CONCERN ASSESSMENT
    # ======================================================================
    add_heading2(doc, "4. Going Concern Assessment")

    add_mixed_para(doc, [
        ("GREEN", True, False, GREEN, Pt(16)),
        (" \u2014 No immediate going concern risk. ", False, False, None, BODY_SIZE),
    ])

    add_para(
        doc,
        "Cash remains positive through 2032 (891.75M). Balance sheet balanced across all "
        "16 periods. Revenue declines in 2031\u20132032 (220\u2192150M) but sufficient cash "
        "reserves from 2026\u20132028 production phase."
    )

    # ======================================================================
    # 5. DD READINESS ASSESSMENT
    # ======================================================================
    add_heading2(doc, "5. DD Readiness Assessment")

    dd_headers = ["Category", "Status", "Key Issue"]
    dd_rows = [
        ["Model Integrity", "FAIL", "Static (97.5% constants)"],
        ["Leakage Control", "FAIL", "23 leakage findings"],
        ["Financial Math", "CONDITIONAL", "Anchors verified but back-solved"],
        ["Valuation", "FAIL", "6\u00d7 spread, fabricated comps"],
        ["MC Simulation", "CONDITIONAL", "Results verified but underpowered at N=5000"],
        ["Documentation", "PASS", "34-item self-audit, 8 ADRs"],
        ["Pipeline", "PASS", "16K LOC, 287 tests, deterministic"],
        ["Cover Docs", "CONDITIONAL", "Honest but incomplete"],
        ["Stress Testing", "CONDITIONAL", "Base case passes, tails weak"],
        ["Legal/Disclaimers", "CONDITIONAL", "Present but template placeholders remain"],
        ["Metadata", "FAIL", "Developer path, internal classification"],
        ["Version Control", "PASS", "Manifest, SHA-256, 12 backup chain"],
    ]
    tbl = make_table(doc, dd_headers, dd_rows, col_widths=[4.5, 3.5, 9.0])

    # Color-code the Status column
    status_colors = {
        "FAIL": "CC0000",
        "CONDITIONAL": "FF8C00",
        "PASS": "228B22",
    }
    for r_idx, row_data in enumerate(dd_rows):
        cell = tbl.rows[r_idx + 1].cells[1]
        status_val = row_data[1]
        if status_val in status_colors:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(status_colors[status_val])
                    run.bold = True

    # ======================================================================
    # 6. ROADMAP TO FULL PASS
    # ======================================================================
    add_heading2(doc, "6. Roadmap to FULL PASS")

    roadmap = [
        ("Quick Wins (0\u20132 days): ",
         "Scrub leakage (items 1\u20132), clean metadata, fix template placeholders, "
         "reconcile scenario probabilities."),
        ("Structural (1\u20133 weeks): ",
         "Convert 5 key sheets to live formulas, standardize IRR method, reconcile "
         "valuation spread to \u22642\u00d7, add comp sources, fix D&A jump."),
        ("Architectural (4\u20136 weeks): ",
         "Full formula-based rebuild, independent verification by second engineer, "
         "MC engine upgrade with proper cashflow-based IRR, bottom-up cost derivation."),
    ]

    for prefix, detail in roadmap:
        p = doc.add_paragraph()
        run_b = p.add_run(prefix)
        run_b.bold = True
        run_b.font.name = BODY_FONT
        run_b._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_b.font.size = BODY_SIZE
        run_b.font.color.rgb = H1_COLOR

        run_d = p.add_run(detail)
        run_d.font.name = BODY_FONT
        run_d._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run_d.font.size = BODY_SIZE
        p.paragraph_format.space_after = Pt(6)

    # ======================================================================
    # 7. P5 VERIFICATION RESULTS
    # ======================================================================
    add_heading2(doc, "7. \u041f5 Verification Results")

    add_para(
        doc,
        "32/32 mechanisms applied across 7 blocks. Confidence level: HIGH for data "
        "extraction and structural checks; MEDIUM for financial logic verification "
        "(limited by static nature of model); HIGH for pipeline assessment."
    )

    # ======================================================================
    # FOOTER / DISCLAIMER
    # ======================================================================
    doc.add_paragraph()  # spacer
    p_hr2 = doc.add_paragraph()
    pPr2 = p_hr2._p.get_or_add_pPr()
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr2.append(pBdr2)

    add_para(
        doc,
        "This document was generated as part of a DD-grade independent audit. "
        "All findings are based on structural analysis of the xlsx artifact, "
        "pipeline code review, and Monte Carlo simulation verification. "
        "This summary does not constitute investment advice.",
        italic=True,
        size=Pt(10),
        color=RGBColor(0x66, 0x66, 0x66),
    )

    add_para(
        doc,
        "Audit Executive Summary v1.0 | April 13, 2026 | Claude Code Auditor | "
        "Confidential \u2014 For Investor Due Diligence Use Only",
        bold=True,
        size=Pt(10),
        color=RGBColor(0x66, 0x66, 0x66),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # -- Save --
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "audit_public_v1_executive_summary.docx")
    doc.save(out_path)
    print(f"[OK] Executive Summary saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    generate()
