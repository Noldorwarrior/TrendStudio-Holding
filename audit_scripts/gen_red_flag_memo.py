#!/usr/bin/env python3
"""
Generate RED FLAG MEMO — DD-grade audit document for TrendStudio investor model.
Output: audit_public_v1_RED_FLAG_MEMO.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def set_run_font(run, name="Times New Roman", size=14, bold=False, italic=False, color=None):
    """Configure font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom H1 styling: 22pt, #0070C0, Times New Roman."""
    heading = doc.add_heading(level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    run.font.name = name = "Times New Roman"
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    else:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    run.font.bold = True
    # Set paragraph spacing tight
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(4)
    return heading


def add_body_paragraph(doc, text="", bold=False, italic=False, alignment=None, space_after=Pt(4)):
    """Add a body paragraph in Times New Roman 14pt."""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = Pt(1)
    if text:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    return para


def add_red_flag_item(doc, number, label, severity, description):
    """Add a formatted red flag item with bold label and severity tag."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.left_indent = Cm(0.5)

    # Number + label (bold)
    run_num = para.add_run(f"{number}. {label}")
    set_run_font(run_num, size=13, bold=True)

    # Severity tag (bold, colored)
    severity_colors = {
        "CRITICAL P0": (0xC0, 0x00, 0x00),
        "HIGH P1": (0xE0, 0x6C, 0x00),
        "MEDIUM P2": (0xBF, 0x8F, 0x00),
    }
    color = severity_colors.get(severity, (0x00, 0x00, 0x00))
    run_sev = para.add_run(f" ({severity})")
    set_run_font(run_sev, size=13, bold=True, color=color)

    # Description
    run_desc = para.add_run(f" \u2014 {description}")
    set_run_font(run_desc, size=13)

    return para


def main():
    doc = Document()

    # --- Page setup: A4, margins 3/1.5/2/2 cm (top/right/bottom/left) ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)

    # --- Set default document font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # ===== TITLE =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run("RED FLAG MEMO")
    set_run_font(run, size=26, bold=True, color=(0xC0, 0x00, 0x00))

    subtitle1 = doc.add_paragraph()
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle1.paragraph_format.space_after = Pt(2)
    run = subtitle1.add_run("Investor Model v1.0.2 Public")
    set_run_font(run, size=18, bold=True, color=(0x00, 0x70, 0xC0))

    # Confidentiality notice
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_after = Pt(2)
    run = conf_para.add_run("CONFIDENTIAL \u2014 DO NOT DISTRIBUTE WITHOUT REMEDIATION")
    set_run_font(run, size=12, bold=True, italic=True, color=(0xC0, 0x00, 0x00))

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(8)
    run = date_para.add_run("April 13, 2026")
    set_run_font(run, size=12)

    # Thin line separator
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(4)
    run = sep.add_run("\u2500" * 72)
    set_run_font(run, size=8, color=(0x99, 0x99, 0x99))

    # ===== SECTION: Critical Red Flags =====
    add_heading_styled(doc, "Critical Red Flags", level=1)

    # --- 10 Red Flag Items ---
    red_flags = [
        (
            1,
            "INTERNAL DATA LEAKAGE",
            "CRITICAL P0",
            'Cell 24_Investor_Returns!B49 explicitly references Internal W5 V\u2011D waterfall: '
            '"Margin to hurdle under W3 (2.09pp) is lower than under Internal W5 V\u2011D (6.75pp)". '
            '8 additional "Internal" keyword occurrences in sharedStrings.xml. W\u2085 unicode subscript present. '
            'IMMEDIATE ACTION: Scrub all "Internal" references before any distribution.'
        ),
        (
            2,
            "DEVELOPER PATH LEAKED",
            "CRITICAL P0",
            "xl/workbook.xml absPath reveals /Users/noldorwarrior/Documents/Claude/Projects/Holding/"
            "Investor_Package/ \u2014 exposes developer username, tool usage (Claude), and internal project "
            "structure. IMMEDIATE ACTION: Re-save xlsx from clean environment or strip x15ac:absPath element."
        ),
        (
            3,
            "STATIC MODEL (97.5% constants)",
            "HIGH P1",
            "Only 226 formulas across 42 sheets (9,000+ value cells). Model is formatted spreadsheet report, "
            "not live financial model. Any analyst attempting sensitivity analysis will discover model doesn\u2019t "
            "respond to input changes. IMMEDIATE ACTION: Convert top-5 critical sheets (P&L, CF, Valuation, MC, "
            "Waterfall) to live formulas before DD."
        ),
        (
            4,
            "ANCHOR-BASED BACK-SOLVING VISIBLE",
            "HIGH P1",
            "Revenue(4,545) = COGS(2,008.8) + OpEx(368.8) + EBITDA(2,167.4) to exact decimal. Decimal precision "
            "reveals fitting, not bottom-up derivation. Known as FIT-01 in self-audit. IMMEDIATE ACTION: Document "
            "as intentional architectural choice; prepare defense for DD."
        ),
        (
            5,
            "VALUATION SPREAD 6\u00d7",
            "HIGH P1",
            "DCF EV = 1.8B, Comps Median EV = 7.5B, MC P50 EV = 11.2B. Gap is 6.2\u00d7 between floor and "
            "ceiling. Industry norm \u22642\u00d7. IMMEDIATE ACTION: Reconcile methodologies to common "
            "horizon/WACC base or explicitly position as floor/fair/ceiling with formal justification."
        ),
        (
            6,
            "MC P(IRR > hurdle) = 13.6%",
            "HIGH P1",
            "Monte Carlo shows only 13.6% probability of beating 18% hurdle rate. Mean MC IRR = 11.44% vs "
            "deterministic Base 20.09% \u2014 gap of 8.65pp. P(Loss) = 5.5%. A sophisticated LP will focus on "
            "MC over deterministic. IMMEDIATE ACTION: Ensure prominent disclosure; prepare narrative explaining "
            "structural reasons (hit_rate, frontload pattern)."
        ),
        (
            7,
            "THREE INCONSISTENT IRR METHODS",
            "HIGH P1",
            "Newton\u2019s method (build_A11), MOIC^(1/6.5)-1 approximation (build_A12), and numpy-based "
            "(v1.0.1 patches) produce different results for same inputs. Base IRR varies from ~7.7% to ~20.09% "
            "depending on method. IMMEDIATE ACTION: Standardize on one verified method (numpy_financial.irr); "
            "document methodology."
        ),
        (
            8,
            "METADATA LEAKS",
            "MEDIUM P2",
            'Description contains anchor "cumulative EBITDA 2026\u20132028 = 3,000 mln RUB", Keywords contain '
            '"L3" (internal classification), Last Modified By = "a" (truncated username), email '
            "team@trendstudio.ru in cells. IMMEDIATE ACTION: Clean metadata before distribution."
        ),
        (
            9,
            "SCENARIO PROBABILITY INCONSISTENCY",
            "MEDIUM P2",
            "build_A12 uses probabilities 5/15/50/20/10, manifest states 10/20/40/20/10. This affects expected "
            "NPV calculations. IMMEDIATE ACTION: Reconcile and use single probability set."
        ),
        (
            10,
            "PEER COMPS FABRICATION RISK",
            "MEDIUM P2",
            "6 peer companies show suspiciously uniform EV/Revenue multiples (0.80\u20130.81 for 5 of 6). No "
            "source attribution. DD valuation team will request Pitchbook/Mergermarket verification. IMMEDIATE "
            "ACTION: Add source citations or flag as indicative."
        ),
    ]

    for num, label, severity, desc in red_flags:
        add_red_flag_item(doc, num, label, severity, desc)

    # ===== SECTION: Verdict =====
    add_heading_styled(doc, "Verdict", level=1)

    verdict_para = doc.add_paragraph()
    verdict_para.paragraph_format.space_after = Pt(4)
    run_v = verdict_para.add_run("FAIL \u2014 DO NOT DISTRIBUTE")
    set_run_font(run_v, size=14, bold=True, color=(0xC0, 0x00, 0x00))
    run_v2 = verdict_para.add_run(
        " in current state. Minimum remediation: fix items 1\u20132 (leakage) before any sharing. "
        "Items 3\u20137 required before formal DD. Items 8\u201310 before CIM."
    )
    set_run_font(run_v2, size=13)

    # ===== SECTION: Effort Estimate =====
    add_heading_styled(doc, "Effort Estimate to CONDITIONAL PASS", level=1)

    effort_items = [
        ("Leakage scrub:", "2\u20134 hours"),
        ("Live formulas conversion:", "2\u20133 weeks"),
        ("IRR standardization:", "1\u20132 days"),
        ("Metadata cleanup:", "1 hour"),
        ("Full remediation to PASS:", "4\u20136 weeks"),
    ]

    for label, estimate in effort_items:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Cm(0.5)
        run_l = para.add_run(f"\u2022  {label} ")
        set_run_font(run_l, size=13, bold=True)
        run_e = para.add_run(estimate)
        set_run_font(run_e, size=13)

    # ===== Save =====
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "audit_public_v1_RED_FLAG_MEMO.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")
    print(f"File size: {os.path.getsize(output_path):,} bytes")


if __name__ == "__main__":
    main()
