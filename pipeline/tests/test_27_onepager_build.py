"""
test_27_onepager_build.py
Тесты для B4 — one-pager executive summary v1.4.3.

Проверяет:
* наличие файла, количество параграфов, таблиц
* ключевые метрики (якорь, LHS mean, VaR, breach, P(reach))
* согласованность со stress_matrix артефактами
* формат: A4, Times New Roman 14pt, поля 2/2/3/1.5
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt

PIPELINE = Path(__file__).resolve().parent.parent
STRESS = PIPELINE / "artifacts" / "stress_matrix"
ARTIFACTS = PIPELINE / "artifacts"

DOCX_PATH = ARTIFACTS / "B4_onepager.docx"


@pytest.fixture(scope="module")
def doc():
    assert DOCX_PATH.exists(), f"Missing {DOCX_PATH}"
    return Document(str(DOCX_PATH))


@pytest.fixture(scope="module")
def lhs_data():
    return json.loads((STRESS / "lhs_copula.json").read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def mc_data():
    return json.loads((STRESS / "monte_carlo.json").read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def gate_data():
    return json.loads((STRESS / "stage_gate.json").read_text(encoding="utf-8"))


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------- structure ----------
def test_onepager_exists(doc):
    assert doc is not None


def test_onepager_has_paragraphs(doc):
    assert len(doc.paragraphs) >= 5


def test_onepager_has_tables(doc):
    """3 таблицы: ключевые показатели, сравнение движков, риски."""
    assert len(doc.tables) >= 3


def test_onepager_has_title(doc):
    txt = _all_text(doc)
    assert "ТрендСтудио" in txt and "2026" in txt and "2028" in txt


def test_onepager_mentions_board_of_directors(doc):
    assert "Совет" in _all_text(doc)


def test_onepager_mentions_version(doc):
    assert "v1.4.3" in _all_text(doc)


# ---------- content ----------
def test_onepager_anchor_3000(doc):
    assert "3 000" in _all_text(doc) or "3000" in _all_text(doc)


def test_onepager_mentions_anchor_tolerance(doc):
    txt = _all_text(doc)
    assert "±1%" in txt or "± 1%" in txt


def test_onepager_lhs_mean(doc, lhs_data):
    mean_int = int(round(lhs_data["mean_ebitda"]))
    expected = f"{mean_int:,}".replace(",", " ")
    assert expected in _all_text(doc)


def test_onepager_var95(doc, lhs_data):
    v = int(round(lhs_data["var_95_mln_rub"]))
    expected = f"{v:,}".replace(",", " ")
    assert expected in _all_text(doc)


def test_onepager_var99(doc, lhs_data):
    v = int(round(lhs_data["var_99_mln_rub"]))
    expected = f"{v:,}".replace(",", " ")
    assert expected in _all_text(doc)


def test_onepager_breach_probability(doc, lhs_data):
    pct_str = f"{lhs_data['breach_probability'] * 100:.2f}".replace(".", ",")
    assert pct_str in _all_text(doc)


def test_onepager_p_reach(doc, gate_data):
    pct_str = f"{gate_data['p_reach_release'] * 100:.2f}".replace(".", ",")
    assert pct_str in _all_text(doc)


def test_onepager_mentions_all_four_engines(doc):
    txt = _all_text(doc)
    assert "Naive MC" in txt
    assert "LHS" in txt and "Copula" in txt
    assert "Bootstrap" in txt
    assert "Stage-Gate" in txt or "Stage-gate" in txt


def test_onepager_has_risks_section(doc):
    txt = _all_text(doc)
    assert "риск" in txt.lower()


def test_onepager_has_mitigation_section(doc):
    assert "Митигация" in _all_text(doc) or "митигация" in _all_text(doc).lower()


def test_onepager_has_recommendation(doc):
    assert "Рекомендация" in _all_text(doc) or "Утвердить" in _all_text(doc)


def test_onepager_fx_mentioned(doc):
    assert "FX" in _all_text(doc)


def test_onepager_stage_gate_mentioned(doc):
    txt = _all_text(doc)
    assert "stage-gate" in txt.lower() or "Stage-Gate" in txt


def test_onepager_mc_breach_in_table(doc, mc_data):
    pct_str = f"{mc_data['breach_probability'] * 100:.2f}".replace(".", ",")
    assert pct_str in _all_text(doc)


# ---------- format ----------
def test_onepager_is_a4_portrait(doc):
    section = doc.sections[0]
    # A4: ~21cm x ~29.7cm
    width_cm = section.page_width / 914400 * 2.54 if False else section.page_width
    # python-docx stores page sizes in EMU
    a4_w = Cm(21)
    a4_h = Cm(29.7)
    tol = Cm(0.2)
    assert abs(section.page_width - a4_w) < tol
    assert abs(section.page_height - a4_h) < tol


def test_onepager_margins(doc):
    section = doc.sections[0]
    tol_emu = 40000  # ~1mm
    assert abs(section.top_margin - Cm(2)) < tol_emu
    assert abs(section.bottom_margin - Cm(2)) < tol_emu
    assert abs(section.left_margin - Cm(3)) < tol_emu
    assert abs(section.right_margin - Cm(1.5)) < tol_emu


def test_onepager_uses_times_new_roman(doc):
    style = doc.styles["Normal"]
    assert style.font.name == "Times New Roman"


def test_onepager_base_font_14pt(doc):
    style = doc.styles["Normal"]
    assert style.font.size == Pt(14)


def test_onepager_volume_reasonable(doc):
    """One-pager должен быть компактным: ≤ 2 страниц A4 ~ ≤ 50 параграфов."""
    total_paragraphs = len(doc.paragraphs)
    assert 5 <= total_paragraphs <= 80, f"Unexpected paragraph count: {total_paragraphs}"
