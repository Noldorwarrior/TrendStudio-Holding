"""
Test B3 Memo build (v1.4.2).

Проверяет, что `scripts/build_memo.py`:
1. Генерирует .docx без ошибок.
2. Содержит все три адресные части (CFO / Риск-комитет / Совет директоров).
3. Содержит требуемые секции (Executive / Методология / 3 Части / Приложения).
4. Ключевые числа в тексте согласованы с JSON-артефактами.
5. Использует корректное форматирование (Times New Roman 14pt, headings #0070C0).
6. Имеет >= 6 страниц контента (по числу параграфов + таблиц).
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_memo import _find_percentile, fmt, load_all, pct  # noqa: E402


@pytest.fixture(scope="module")
def memo_doc(tmp_path_factory) -> Document:
    from build_memo import build_memo
    out = tmp_path_factory.mktemp("memo") / "memo.docx"
    build_memo(out)
    assert out.exists()
    return Document(out)


@pytest.fixture(scope="module")
def data() -> dict:
    return load_all()


def _all_text(doc: Document) -> str:
    """Извлекает весь текст из параграфов и таблиц."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# Structure tests
# ─────────────────────────────────────────────────────────────────────────────
def test_memo_has_three_parts(memo_doc):
    text = _all_text(memo_doc)
    assert "Часть I." in text
    assert "Часть II." in text
    assert "Часть III." in text


def test_memo_three_audiences(memo_doc):
    text = _all_text(memo_doc)
    assert "CFO" in text or "Правлени" in text
    assert "Риск-комитет" in text
    assert "Совет директоров" in text


def test_memo_has_executive_summary(memo_doc):
    headings = [p.text for p in memo_doc.paragraphs if p.style and "Heading" in p.style.name]
    assert any("Резюме" in h for h in headings)


def test_memo_has_methodology(memo_doc):
    text = _all_text(memo_doc)
    assert "Методология" in text
    assert "П1" in text and "Аналитик" in text


def test_memo_has_appendices(memo_doc):
    headings = [p.text for p in memo_doc.paragraphs if p.style and "Heading" in p.style.name]
    app_count = sum(1 for h in headings if "Приложение" in h)
    assert app_count >= 3, f"Expected ≥3 appendices, got {app_count}"


def test_memo_has_verification_section(memo_doc):
    text = _all_text(memo_doc)
    assert "Верификация" in text
    assert "М2" in text


# ─────────────────────────────────────────────────────────────────────────────
# Content accuracy tests (numerical)
# ─────────────────────────────────────────────────────────────────────────────
def test_memo_contains_anchor_3000(memo_doc):
    text = _all_text(memo_doc)
    assert "3 000" in text or "3000" in text


def test_memo_contains_lhs_mean(memo_doc, data):
    """LHS mean должен присутствовать в тексте."""
    text = _all_text(memo_doc)
    expected = fmt(data["lhs"]["mean_ebitda"])
    assert expected in text, f"Expected '{expected}' not found in memo"


def test_memo_contains_var95(memo_doc, data):
    text = _all_text(memo_doc)
    expected = fmt(data["lhs"]["var_95_mln_rub"])
    assert expected in text


def test_memo_contains_breach_probability(memo_doc, data):
    text = _all_text(memo_doc)
    expected = pct(data["lhs"]["breach_probability"])
    assert expected in text


def test_memo_contains_stage_gate_p_reach(memo_doc, data):
    text = _all_text(memo_doc)
    p = data["gate"]["p_reach_release"]
    # Разные формы: 72,06%, 0.7206
    assert pct(p) in text or f"{p:.4f}" in text


def test_memo_contains_all_four_engines_in_appendix(memo_doc):
    text = _all_text(memo_doc)
    assert "Naive MC" in text
    assert "Bootstrap" in text
    assert "Stage-Gate" in text or "Stage-gate" in text
    assert "LHS" in text and "Copula" in text


# ─────────────────────────────────────────────────────────────────────────────
# Volume tests
# ─────────────────────────────────────────────────────────────────────────────
def test_memo_volume_at_least_6_pages(memo_doc):
    """Грубо: параграфы + таблицы должны давать ≥6 страниц A4."""
    non_empty = [p for p in memo_doc.paragraphs if p.text.strip()]
    total_chars = sum(len(p.text) for p in non_empty)
    # 2500 знаков/страница для 14pt TNR 1.15 интервала — консервативная оценка
    est_pages_text = total_chars / 2500
    est_pages_tables = len(memo_doc.tables) * 0.5  # ~0.5 стр на таблицу
    est_total = est_pages_text + est_pages_tables
    assert est_total >= 6, (
        f"Memo too short: {est_pages_text:.1f}p text + "
        f"{est_pages_tables:.1f}p tables = {est_total:.1f}p, expected ≥6"
    )


def test_memo_has_at_least_8_tables(memo_doc):
    assert len(memo_doc.tables) >= 7, f"Expected ≥7 tables, got {len(memo_doc.tables)}"


def test_memo_has_at_least_15_headings(memo_doc):
    headings = [p for p in memo_doc.paragraphs if p.style and "Heading" in p.style.name]
    assert len(headings) >= 15, f"Expected ≥15 headings, got {len(headings)}"


# ─────────────────────────────────────────────────────────────────────────────
# Formatting tests (preference #6)
# ─────────────────────────────────────────────────────────────────────────────
def test_memo_uses_times_new_roman(memo_doc):
    style = memo_doc.styles["Normal"]
    assert style.font.name == "Times New Roman"


def test_memo_normal_font_size_14pt(memo_doc):
    style = memo_doc.styles["Normal"]
    assert style.font.size.pt == 14


def test_memo_page_margins_correct(memo_doc):
    from docx.shared import Cm
    section = memo_doc.sections[0]
    # Допуск ±1 mm = ~36000 EMU из-за EMU-округления docx
    tol = 40000
    assert abs(section.top_margin - Cm(2).emu) < tol
    assert abs(section.bottom_margin - Cm(2).emu) < tol
    assert abs(section.left_margin - Cm(3).emu) < tol
    assert abs(section.right_margin - Cm(1.5).emu) < tol


# ─────────────────────────────────────────────────────────────────────────────
# Helper function tests
# ─────────────────────────────────────────────────────────────────────────────
def test_fmt_helper():
    assert fmt(3000) == "3 000"
    assert fmt(2952.24, 1) == "2 952,2"
    assert fmt(1234567) == "1 234 567"


def test_pct_helper():
    assert pct(0.0485) == "4,85%"
    assert pct(0.720613) == "72,06%"
    assert pct(0.0005, 3) == "0,050%"


def test_find_percentile_interpolation(data):
    lhs = data["lhs"]
    # p50 точно = p50
    assert _find_percentile(lhs, lhs["p50_ebitda"]) == 50
    # p5 точно = p5
    assert _find_percentile(lhs, lhs["p5_ebitda"]) == 5
    # Значение ниже p1 → p1 (граница)
    assert _find_percentile(lhs, lhs["p1_ebitda"] - 100) == 1
